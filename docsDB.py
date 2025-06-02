import chromadb
import chromadb.utils.embedding_functions as embedding_functions
import pandas as pd
from docx import Document

chroma_client = chromadb.PersistentClient(path="docs")

openai_ef = embedding_functions.OpenAIEmbeddingFunction(
                #api_key="YOUR_API_KEY",
                model_name="text-embedding-3-large"
            )
collection = chroma_client.get_collection(name="docs", embedding_function=openai_ef)

# Load the DOCX file
file_path = "Additional info (brand, reviews, customer tickets).docx"
doc = Document(file_path)

# Extract brand philosophy text (all paragraphs until first table)
paras = []
for para in doc.paragraphs:
    if para.text.strip() == "":
        continue
    paras.append(para.text)
    # Stop if we detect start of tables in document flow; 
    # but python-docx doesn't link paragraphs to tables directly. 
    # Instead, we'll stop before tables by separate logic below.

# Extract tables
tables = doc.tables

# Extract brand philosophy by taking paragraphs up to a keyword
brand_philosophy = []
for para in doc.paragraphs:
    if para.text.strip().startswith("Reviewer"):
        break
    if para.text.strip():
        brand_philosophy.append(para.text)

brand_philosophy_text = "\n".join(brand_philosophy)

# Process reviews table (first table)
reviews = []
review_metadata = []
review_ids = []
for i, row in enumerate(tables[0].rows):
    if i == 0:
        # header row
        headers = [cell.text.strip() for cell in row.cells]
    else:
        cells = [cell.text.strip() for cell in row.cells]
        data = dict(zip(headers, cells))
        if len(data["Review"])==0:
          continue
        # Create document and metadata
        doc_text = data["Review"] or ""
        metadata = {
            "Reviewer": data["Reviewer"] or "",
            "Product": data["Product"] or "",
            "Rating": data["Rating"] or ""
        }
        reviews.append(doc_text)
        review_metadata.append(metadata)
        review_ids.append(f"review_{i}")

# Process customer tickets table (second table)
tickets = []
ticket_metadata = []
ticket_ids = []
for i, row in enumerate(tables[1].rows):
    if i == 0:
        headers = [cell.text.strip() for cell in row.cells]
    else:
        cells = [cell.text.strip() for cell in row.cells]
        data = dict(zip(headers, cells))
        doc_text = f"Customer Message: {data['Customer Message']}. Support Response: {data['Support Response']}"
        metadata = {
            "TicketID": data["Ticket ID"] or ""
        }
        tickets.append(doc_text)
        ticket_metadata.append(metadata)
        ticket_ids.append(data["Ticket ID"])

# Prepare brand philosophy as separate entry
bp_document = brand_philosophy_text
bp_metadata = {"section": "Brand Philosophy"}
bp_id = "brand_philosophy"

# Combine all
documents = [bp_document] + reviews + tickets
metadatas = [bp_metadata] + review_metadata + ticket_metadata
ids = [bp_id] + review_ids + ticket_ids

for i in range(len(documents)):
  collection.add(
      documents=[documents[i]],
      metadatas=[metadatas[i]],
      ids=[ids[i]]
  )